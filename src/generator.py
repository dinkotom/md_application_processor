import qrcode
from docx import Document
from docx.shared import Inches
from typing import Dict
import os

def generate_qr_code(data: Dict[str, str], output_path: str) -> str:
    """
    Generates a QR code image from the application data.
    Payload format: First|Last|DOB|Email|MembershipID
    """
    # Construct payload
    payload = (
        f"Jméno: {data.get('first_name', '')}\n"
        f"Příjmení: {data.get('last_name', '')}\n"
        f"Číslo karty: {data.get('membership_id', '')}\n"
        f"Datum narození: {data.get('dob', '')}\n"
        f"Email: {data.get('email', '')}\n"
        f"Telefon: {data.get('phone', '')}"
    )
    
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(payload)
    qr.make(fit=True)

    img = qr.make_image(fill_color="#ffffff", back_color="#36378c")
    img.save(output_path)
    return output_path

def generate_qr_code_bytes(data: Dict[str, str]):
    """
    Generates a QR code image from the application data and returns it as bytes.
    Payload format: First|Last|DOB|Email|MembershipID
    """
    from io import BytesIO
    
    # Construct payload
    payload = (
        f"Jméno: {data.get('first_name', '')}\n"
        f"Příjmení: {data.get('last_name', '')}\n"
        f"Číslo karty: {data.get('membership_id', '')}\n"
        f"Datum narození: {data.get('dob', '')}\n"
        f"Email: {data.get('email', '')}\n"
        f"Telefon: {data.get('phone', '')}"
    )
    
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(payload)
    qr.make(fit=True)

    img = qr.make_image(fill_color="#ffffff", back_color="#36378c")
    
    # Save to BytesIO instead of file
    img_io = BytesIO()
    img.save(img_io, 'PNG')
    img_io.seek(0)
    return img_io

from datetime import datetime

def calculate_age(dob_str: str) -> int:
    """Calculates age from DOB string. Supports DD/MM/YYYY and DD.MM.YYYY."""
    try:
        # Normalize separators
        clean_dob = dob_str.strip().replace(" ", "").replace(".", "/")
        # Handle cases like 14/5/2000 -> 14/05/2000
        parts = clean_dob.split("/")
        if len(parts) == 3:
            clean_dob = f"{int(parts[0]):02d}/{int(parts[1]):02d}/{parts[2]}"
            
        dob = datetime.strptime(clean_dob, "%d/%m/%Y")
        today = datetime.today()
        age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
        return age
    except ValueError:
        return 0

import unicodedata

def normalize_text(text: str) -> str:
    """Removes accents and converts to lowercase."""
    if not text:
        return ""
    nfkd_form = unicodedata.normalize('NFKD', text)
    return "".join([c for c in nfkd_form if not unicodedata.combining(c)]).lower()

def generate_document(data: Dict[str, str], qr_code_path: str, output_path: str, is_duplicate: bool = False):
    """
    Generates a Word document with the application details and QR code.
    """
    doc = Document()
    
    # 1. Header: Name Surname
    full_name = f"{data.get('first_name', '')} {data.get('last_name', '')}"
    doc.add_heading(full_name, 0)
    
    # 2. Details
    doc.add_paragraph(f"Číslo karty: {data.get('membership_id', '')}")
    doc.add_paragraph(f"Datum narození: {data.get('dob', '')}")
    doc.add_paragraph(f"Email: {data.get('email', '')}")
    doc.add_paragraph(f"Telefon: {data.get('phone', '')}")
    
    # 3. Warnings
    warnings = []
    
    # Age Check
    age = calculate_age(data.get('dob', ''))
    if age < 15 or age >= 25:
        warnings.append(f"⚠️ POZOR: Uchazeč má {age} let.")
        
    # Email-Name Check
    last_name = data.get('last_name', '')
    email = data.get('email', '')
    
    norm_last = normalize_text(last_name)
    norm_email = normalize_text(email)
    
    if norm_last and norm_last not in norm_email:
        warnings.append("⚠️ POZOR: Email neodpovídá jménu.")
        
    # Duplicate Check
    if is_duplicate:
        warnings.append("⚠️ POZOR: Duplicitní přihláška.")
        
    if warnings:
        doc.add_paragraph()
        for w in warnings:
            run = doc.add_paragraph().add_run(w)
            run.bold = True
    
    # 4. QR Code (Centered)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = 1 # Center
    run = p.add_run()
    run.add_picture(qr_code_path, width=Inches(2.0))
    
    # 5. Original Email
    doc.add_heading('Původní e-mail:', level=2)
    doc.add_paragraph(data.get('full_body', ''))
    
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    # Test run
    test_data = {
        'first_name': 'Barbora',
        'last_name': 'Smékalová',
        'email': 'barus.smekalova@outlook.cz',
        'dob': '14/05/2000',
        'membership_id': '1954',
        'full_body': 'Full text content here...'
    }
    
    # Create a temp dir for artifacts if needed, or just use current
    qr_path = "test_qr.png"
    doc_path = "test_application.docx"
    
    generate_qr_code(test_data, qr_path)
    generate_document(test_data, qr_path, doc_path)
    
    # Cleanup
    if os.path.exists(qr_path):
        os.remove(qr_path)
